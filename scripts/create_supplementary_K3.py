#!/usr/bin/env python3
"""ModelIt! K-3 Supplementary Materials Generator (TPT-Quality Visual Edition)

Generates visually rich, kid-friendly supplementary materials:
  K-2: Model Explorer Game (HTML), Coloring Sheet (DOCX - landscape, full-page art)
  G2-3: Fix the System (DOCX - landscape, visual), Interactive Notebook (DOCX),
        Reading Passage (DOCX - picture-driven), Parent-Home Connection (DOCX)
  Grade 2 gets ALL 6 types.

Design philosophy:
  K-1: VERY visual, minimal text, huge illustrations, TPT coloring page style
  G2-3: Still picture-driven, tighter worksheet format, landscape fix sheets

Usage:
    python create_supplementary_K3.py K      # Kindergarten: game + coloring
    python create_supplementary_K3.py 1      # Grade 1: game + coloring
    python create_supplementary_K3.py 2      # Grade 2: all 6 types
    python create_supplementary_K3.py 3      # Grade 3: fix + notebook + reading + home
    python create_supplementary_K3.py all    # All grades
"""
import os, sys, time, base64, requests, importlib, json, random
from docx import Document
from docx.shared import Inches as DI, Pt as DP, RGBColor as DR, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import dotenv as _dotenv
_dotenv.load_dotenv(os.path.join(os.path.expanduser('~'), '.env'))
API_KEY = os.environ.get('OPENROUTER_API_KEY', '')
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Brand colors ──
DN  = DR(0x0D, 0x1B, 0x2A)
DBB = DR(0x1A, 0x47, 0x80)
DMB = DR(0x2E, 0x86, 0xAB)
DLB = DR(0x7E, 0xC8, 0xE3)
DOG = DR(0xE6, 0x7E, 0x22)
DWH = DR(0xFF, 0xFF, 0xFF)
DGN = DR(0x27, 0xAE, 0x60)
DRD = DR(0xE7, 0x4C, 0x3C)
DPURP = DR(0x9B, 0x59, 0xB6)
DPINK = DR(0xE9, 0x1E, 0x63)
DYEL  = DR(0xF3, 0x9C, 0x12)

HEX_NAVY  = "0D1B2A"; HEX_BB = "1A4780"; HEX_MB = "2E86AB"
HEX_LB    = "7EC8E3"; HEX_SB = "5DB7DE"; HEX_OG = "E67E22"
HEX_GREEN = "27AE60"; HEX_RED = "E74C3C"; HEX_LGRAY = "F5F5F5"
HEX_YELLOW = "FFF9C4"; HEX_LGREEN = "E8F5E9"; HEX_LPINK = "FCE4EC"
HEX_LBLUE = "E3F2FD"; HEX_LPURP = "F3E5F5"; HEX_LORANGE = "FFF3E0"

total_cost = 0.0

GRADE_CONFIG = {
    "K":  {"label": "Kindergarten", "age": "5-6 years old", "tier": 1,
           "dir": "grade-K",  "data": "lesson_data_GK",
           "game": True, "coloring": True, "fix": False, "notebook": False,
           "reading": False, "home": False},
    "1":  {"label": "1st Grade",    "age": "6-7 years old", "tier": 1,
           "dir": "grade-01", "data": "lesson_data_G01",
           "game": True, "coloring": True, "fix": False, "notebook": False,
           "reading": False, "home": False},
    "2":  {"label": "2nd Grade",    "age": "7-8 years old", "tier": 2,
           "dir": "grade-02", "data": "lesson_data_G02",
           "game": True, "coloring": True, "fix": True, "notebook": True,
           "reading": True, "home": True},
    "3":  {"label": "3rd Grade",    "age": "8-9 years old", "tier": 2,
           "dir": "grade-03", "data": "lesson_data_G03",
           "game": False, "coloring": False, "fix": True, "notebook": True,
           "reading": True, "home": True},
}


# ================================================================
#  SHARED UTILITIES
# ================================================================

def cell_shd(cell, hex_c):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), hex_c)
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_border(cell, **kwargs):
    """Set cell borders. e.g. set_cell_border(cell, top=("sz","12","color","000000","val","single"))"""
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, attrs in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        for i in range(0, len(attrs), 2):
            element.set(qn(f'w:{attrs[i]}'), attrs[i+1])
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_width(cell, width_inches):
    """Force a cell to a specific width."""
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_row_height(row, height_inches):
    """Set exact row height."""
    tr = row._tr; trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_inches * 1440)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def remove_table_borders(table):
    """Remove all borders from a table for invisible layout."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    tblPr.append(borders)

def make_landscape(doc):
    """Set document to landscape with tight margins."""
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = DI(11)
        section.page_height = DI(8.5)
        section.top_margin = DI(0.3)
        section.bottom_margin = DI(0.3)
        section.left_margin = DI(0.4)
        section.right_margin = DI(0.4)

def make_portrait_tight(doc):
    """Set document to portrait with tight margins."""
    for section in doc.sections:
        section.top_margin = DI(0.35)
        section.bottom_margin = DI(0.35)
        section.left_margin = DI(0.5)
        section.right_margin = DI(0.5)

def add_doodle_border_table(doc, bg_hex="FFFFFF"):
    """Create a full-page rounded-feel border using a table with colored thick borders."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell_shd(cell, bg_hex)
    # Thick decorative border
    for edge in ['top', 'left', 'bottom', 'right']:
        set_cell_border(cell, **{edge: ("sz", "24", "color", HEX_MB, "val", "single")})
    return cell

def fun_header(cell_or_doc, title, subtitle="", font_size=28, color=DMB):
    """Add a large fun header suitable for K-3."""
    p = cell_or_doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True; r.font.size = DP(font_size); r.font.color.rgb = color
    r.font.name = "Comic Sans MS"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = DP(2)
    if subtitle:
        p2 = cell_or_doc.add_paragraph()
        r2 = p2.add_run(subtitle)
        r2.font.size = DP(14); r2.font.color.rgb = DR(0x55, 0x55, 0x55)
        r2.font.name = "Comic Sans MS"
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = DP(2)
    return p

def name_date_line(cell_or_doc, font_size=14):
    """Kid-friendly name/date line."""
    p = cell_or_doc.add_paragraph()
    r = p.add_run("Name: ________________________    Date: ____________")
    r.font.size = DP(font_size); r.font.name = "Comic Sans MS"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = DP(4)
    p.paragraph_format.space_after = DP(6)

def writing_lines(cell_or_doc, count=3, width=50):
    """Wide-spaced dotted writing lines for young writers."""
    for _ in range(count):
        p = cell_or_doc.add_paragraph()
        r = p.add_run("_" * width)
        r.font.size = DP(16); r.font.color.rgb = DR(0xBB, 0xBB, 0xBB)
        r.font.name = "Comic Sans MS"
        p.paragraph_format.space_after = DP(2)
        p.paragraph_format.line_spacing = DP(28)

def compose_sentence_frame(c):
    if "sentence_frame" in c and c["sentence_frame"]:
        return c["sentence_frame"]
    rels = c.get("relationships", [])
    if rels:
        rel_str, sign, _ = rels[0]
        parts = rel_str.split(" to ")
        if len(parts) == 2:
            fr, to = parts[0].strip(), parts[1].strip()
            return f"When {fr} goes up, {to} goes _______."
    comps = c.get("components", [])
    ext = [n for n, _, e in comps if e]
    int_ = [n for n, _, e in comps if not e]
    if ext and int_:
        return f"When {ext[0]} changes, {int_[0]} _______."
    return "When _______ changes, _______ _______."


# ================================================================
#  AI IMAGE GENERATION
# ================================================================

def gen_img(prompt, fn, img_dir):
    """Generate an image via OpenRouter Gemini Flash Image."""
    global total_cost
    fp = os.path.join(img_dir, fn)
    if os.path.exists(fp):
        print(f"    [SKIP] {fn}"); return fp
    print(f"    Generating: {fn}...")
    try:
        r = requests.post("https://openrouter.ai/api/v1/chat/completions",
            headers={"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"},
            json={"model": "google/gemini-2.5-flash-image",
                  "messages": [{"role": "user", "content": prompt}],
                  "modalities": ["image", "text"]},
            timeout=180)
        if r.status_code == 200:
            data = r.json()
            imgs = data.get("choices", [{}])[0].get("message", {}).get("images", [])
            if imgs:
                img = imgs[0]
                url = img.get("image_url", {}).get("url", "") if isinstance(img, dict) else ""
                if url.startswith("data:image"):
                    with open(fp, "wb") as f:
                        f.write(base64.b64decode(url.split(",")[1]))
                    total_cost += data.get("usage", {}).get("cost", 0)
                    print(f"    [OK] {fn}"); return fp
            print(f"    [!] No image data"); return None
        else:
            print(f"    [X] {r.status_code}"); return None
    except Exception as e:
        print(f"    [X] {e}"); return None


def gen_coloring_img(c, img_dir):
    """Generate TPT-style coloring page image (full-page, thick outlines, landscape)."""
    comps = c.get("components", [])
    comp_names = [n for n, _, _ in comps]
    rels = c.get("relationships", [])
    title = c["title"]

    # Build relationship descriptions
    arrows = []
    for rel_str, sign, _ in rels:
        parts = rel_str.split(" to ")
        if len(parts) == 2:
            d = "together (both go up)" if "+" in sign else "opposite (one up, one down)"
            arrows.append(f"arrow from '{parts[0].strip()}' to '{parts[1].strip()}' going {d}")

    prompt = (
        f"Create a full-page coloring book illustration for young children (ages 5-8) about "
        f"the science topic: '{title}'. "
        f"The scene should show the system model with these parts: {', '.join(comp_names)}. "
        f"Include {' and '.join(arrows) if arrows else 'arrows connecting parts'}. "
        "STYLE: Teachers Pay Teachers quality coloring page. "
        "Thick bold black outlines only, NO colors, NO shading, NO gradients, NO gray tones. "
        "Pure white background with pure black lines. "
        "Large simple cartoon shapes that are EASY to color with crayons. "
        "Each component should be inside a big rounded bubble or fun shape with its name "
        "written in large kid-friendly block letters. "
        "Big chunky arrows with + or - symbols between shapes. "
        "Add cute decorative elements around the edges: small stars, swirls, clouds, or "
        "science doodles (beakers, magnifying glasses, light bulbs). "
        "Include a decorative banner at top saying 'Color My Model!' in fun block letters. "
        "Landscape orientation (wider than tall). "
        "Style like a premium TPT coloring worksheet - clean, professional, fun, engaging."
    )
    return gen_img(prompt, f"{c['id']}-coloring.png", img_dir)


def gen_scene_img(c, img_dir, suffix="scene"):
    """Generate a colorful scene illustration for reading passages and home connections."""
    title = c["title"]
    prompt = (
        f"Create a bright, colorful, kid-friendly cartoon illustration about '{title}' "
        "for an elementary science worksheet. "
        "Style: cheerful flat illustration with bright saturated colors, "
        "simple shapes, friendly faces on objects/animals if appropriate. "
        "Clean white background, professional educational illustration quality. "
        "Include cute science-themed decorative elements (stars, question marks, arrows). "
        "Landscape format, suitable for a children's worksheet header. "
        "Age-appropriate for 7-9 year olds. Fun and inviting."
    )
    return gen_img(prompt, f"{c['id']}-{suffix}.png", img_dir)


def gen_fix_img(c, img_dir):
    """Generate a visual model diagram for fix-the-system worksheets."""
    comps = c.get("components", [])
    comp_names = [n for n, _, _ in comps]
    prompt = (
        f"Create a colorful educational diagram for children showing a simple system model. "
        f"The model has {len(comps)} parts: {', '.join(comp_names)}. "
        "Draw each part as a bright colored circle or rounded rectangle with the name "
        "written inside in bold kid-friendly text. "
        "Connect them with thick arrows. Use bright primary colors (blue, green, orange, purple). "
        "Style: clean, modern educational illustration. White background. "
        "Landscape format. Like a Teachers Pay Teachers science diagram. "
        "Add small decorative science doodles in corners (magnifying glass, gears, lightbulb). "
        "Age-appropriate for grades 2-3 (ages 7-9)."
    )
    return gen_img(prompt, f"{c['id']}-fix-diagram.png", img_dir)


# ================================================================
#  1. MODEL EXPLORER GAME (K-2) - HTML
# ================================================================

GAME_HTML_TEMPLATE = r'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Model Explorer: {{TITLE}}</title>
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Comic Sans MS','Segoe UI',system-ui,sans-serif;background:#f0f7ff;
  min-height:100vh;display:flex;flex-direction:column;align-items:center}
.header{background:linear-gradient(135deg,#1A4780,#2E86AB);color:#fff;
  width:100%;padding:16px 24px;text-align:center;box-shadow:0 2px 8px rgba(0,0,0,.15)}
.header h1{font-size:clamp(22px,5vw,36px);margin-bottom:4px}
.header .sub{font-size:clamp(14px,3vw,20px);opacity:.9}
.nav{display:flex;gap:10px;justify-content:center;padding:14px;flex-wrap:wrap}
.nav button{padding:12px 24px;border:4px solid #1A4780;border-radius:30px;
  background:#fff;color:#1A4780;font-size:clamp(15px,3vw,20px);font-weight:700;
  cursor:pointer;transition:all .2s;font-family:inherit}
.nav button:hover,.nav button.active{background:#1A4780;color:#fff;transform:scale(1.08)}
.nav button.done{border-color:#27AE60;color:#27AE60}
.nav button.done.active{background:#27AE60;color:#fff}
.screen{display:none;width:min(95vw,750px);padding:20px;margin:8px auto}
.screen.active{display:block}
.screen h2{color:#1A4780;font-size:clamp(22px,4vw,30px);text-align:center;margin-bottom:16px}
.instruction{background:#fff;border-left:6px solid #E67E22;padding:14px 18px;
  border-radius:0 12px 12px 0;margin-bottom:20px;font-size:clamp(16px,3vw,20px);
  color:#333;box-shadow:0 2px 6px rgba(0,0,0,.06)}

/* Sort */
.sort-area{display:flex;gap:20px;justify-content:center;flex-wrap:wrap;margin-bottom:20px}
.sort-box{background:#fff;border:4px dashed #2E86AB;border-radius:20px;padding:18px;
  min-width:220px;min-height:160px;flex:1;text-align:center;transition:all .3s}
.sort-box h3{color:#fff;padding:10px;border-radius:12px;margin-bottom:14px;
  font-size:clamp(15px,3vw,20px)}
.sort-box.external h3{background:linear-gradient(135deg,#E67E22,#F39C12)}
.sort-box.internal h3{background:linear-gradient(135deg,#2E86AB,#5DB7DE)}
.sort-box.highlight{border-color:#E67E22;background:#fff8f0;transform:scale(1.02)}
.cards{display:flex;gap:12px;justify-content:center;flex-wrap:wrap;margin-bottom:18px}
.card{background:linear-gradient(135deg,#5DB7DE,#2E86AB);color:#fff;
  padding:14px 24px;border-radius:16px;font-size:clamp(16px,3vw,22px);
  font-weight:700;cursor:grab;user-select:none;touch-action:none;
  box-shadow:0 4px 10px rgba(0,0,0,.15);transition:transform .2s}
.card:active{transform:scale(1.1);cursor:grabbing}
.card.placed{opacity:.3;pointer-events:none}
.placed-card{background:#E8F5E9;border:3px solid #27AE60;border-radius:12px;
  padding:10px 16px;margin:5px;display:inline-block;font-weight:700;color:#1A4780;
  font-size:clamp(15px,3vw,18px)}

/* Arrows */
.model-display{background:#fff;border-radius:20px;padding:28px;
  box-shadow:0 3px 15px rgba(0,0,0,.08);text-align:center;margin-bottom:18px}
.comp-row{display:flex;align-items:center;justify-content:center;gap:16px;
  flex-wrap:wrap;margin:16px 0}
.comp-box{background:linear-gradient(135deg,#E3F2FD,#BBDEFB);border:4px solid #1A4780;
  border-radius:16px;padding:16px 24px;font-weight:700;font-size:clamp(16px,3vw,22px);
  color:#1A4780;min-width:130px}
.arrow-btn{width:72px;height:72px;border-radius:50%;border:4px solid #ccc;
  font-size:32px;font-weight:900;cursor:pointer;transition:all .3s;
  display:flex;align-items:center;justify-content:center;background:#fff;font-family:inherit}
.arrow-btn.positive{border-color:#27AE60;color:#27AE60;background:#E8F5E9}
.arrow-btn.negative{border-color:#E74C3C;color:#E74C3C;background:#FDEDEC}
.arrow-btn:hover{transform:scale(1.15)}
.arrow-label{font-size:clamp(12px,2.5vw,15px);color:#666;margin-top:6px;text-align:center;font-weight:600}

/* Explore */
.explore-panel{background:#fff;border-radius:20px;padding:28px;
  box-shadow:0 3px 15px rgba(0,0,0,.08);text-align:center}
.ctrl-row{display:flex;align-items:center;justify-content:center;gap:18px;
  margin:18px 0;flex-wrap:wrap}
.ctrl-label{font-size:clamp(16px,3vw,22px);font-weight:700;color:#1A4780;min-width:110px}
.ctrl-btn{width:64px;height:64px;border-radius:50%;border:none;
  font-size:32px;font-weight:900;cursor:pointer;color:#fff;transition:transform .2s;font-family:inherit}
.ctrl-btn:hover{transform:scale(1.15)}
.ctrl-btn.up{background:linear-gradient(135deg,#27AE60,#2ECC71)}
.ctrl-btn.down{background:linear-gradient(135deg,#E74C3C,#C0392B)}
.bar-container{width:min(80%,300px);height:36px;background:#eee;border-radius:18px;
  overflow:hidden;margin:8px auto;position:relative}
.bar-fill{height:100%;border-radius:18px;transition:width .5s ease;
  display:flex;align-items:center;justify-content:flex-end;padding-right:10px;
  font-size:15px;font-weight:700;color:#fff}
.bar-fill.ext{background:linear-gradient(90deg,#E67E22,#F39C12)}
.bar-fill.int{background:linear-gradient(90deg,#2E86AB,#5DB7DE)}
.result-text{font-size:clamp(16px,3vw,20px);color:#333;margin:14px 0;
  min-height:50px;line-height:1.5}
.sentence{background:#FFF8E1;border:3px solid #F39C12;border-radius:16px;
  padding:16px 24px;font-size:clamp(16px,3vw,20px);font-weight:600;
  color:#1A4780;margin-top:18px}

/* Celebration */
.celebrate{position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:999}
.star{position:absolute;font-size:clamp(24px,6vw,44px);animation:fall 1.8s ease-out forwards}
@keyframes fall{0%{opacity:1;transform:translateY(-20px) rotate(0deg) scale(1)}
  100%{opacity:0;transform:translateY(80vh) rotate(720deg) scale(.3)}}
.check-anim{animation:pop .5s ease}
@keyframes pop{0%{transform:scale(0)}50%{transform:scale(1.3)}100%{transform:scale(1)}}
.feedback{padding:14px 24px;border-radius:16px;text-align:center;
  font-size:clamp(16px,3vw,20px);font-weight:700;margin:14px 0;animation:pop .4s ease}
.feedback.correct{background:#E8F5E9;color:#27AE60;border:3px solid #27AE60}
.feedback.wrong{background:#FDEDEC;color:#E74C3C;border:3px solid #E74C3C}
.reset-btn{display:block;margin:20px auto 0;padding:12px 28px;border:3px solid #2E86AB;
  border-radius:24px;background:#fff;color:#2E86AB;font-size:18px;font-weight:700;
  cursor:pointer;font-family:inherit}
.reset-btn:hover{background:#2E86AB;color:#fff}
</style>
</head>
<body>
<div class="header">
  <h1>Model Explorer</h1>
  <div class="sub" id="lessonTitle"></div>
</div>
<div class="nav">
  <button id="nav1" class="active" onclick="showScreen(1)">1. Sort the Parts!</button>
  <button id="nav2" onclick="showScreen(2)">2. Fix the Arrows!</button>
  <button id="nav3" onclick="showScreen(3)">3. What Happens?</button>
</div>
<div class="screen active" id="screen1">
  <h2>Sort the Parts!</h2>
  <div class="instruction">Drag each part into the correct box. <b>We Control</b> means WE
    decide to change it. <b>Changes by Itself</b> means it changes on its own!</div>
  <div class="cards" id="cardArea"></div>
  <div class="sort-area">
    <div class="sort-box external" id="extBox" ondragover="allowDrop(event)"
         ondrop="dropCard(event,'external')">
      <h3>We Control (Outside)</h3><div id="extCards"></div>
    </div>
    <div class="sort-box internal" id="intBox" ondragover="allowDrop(event)"
         ondrop="dropCard(event,'internal')">
      <h3>Changes by Itself (Inside)</h3><div id="intCards"></div>
    </div>
  </div>
  <div id="sortFeedback"></div>
</div>
<div class="screen" id="screen2">
  <h2>Fix the Arrows!</h2>
  <div class="instruction">The arrows are WRONG! Tap each arrow to change it between
    <span style="color:#27AE60;font-weight:700">+ (go together)</span> and
    <span style="color:#E74C3C;font-weight:700">- (go opposite)</span> until correct!</div>
  <div class="model-display" id="arrowModel"></div>
  <div id="arrowFeedback"></div>
</div>
<div class="screen" id="screen3">
  <h2>What Happens?</h2>
  <div class="instruction">Press UP or DOWN buttons to change the outside part.
    Watch what happens to the inside parts!</div>
  <div class="explore-panel" id="explorePanel"></div>
</div>
<script>
const LESSON = {{LESSON_JSON}};
let screenDone=[false,false,false],sortCorrect={},arrowStates={},exploreLevels={};
document.getElementById('lessonTitle').textContent=LESSON.title;
initSort();initArrows();initExplore();

function showScreen(n){
  document.querySelectorAll('.screen').forEach(s=>s.classList.remove('active'));
  document.querySelectorAll('.nav button').forEach(b=>b.classList.remove('active'));
  document.getElementById('screen'+n).classList.add('active');
  document.getElementById('nav'+n).classList.add('active');
}
function initSort(){
  const area=document.getElementById('cardArea');
  let comps=[...LESSON.components];
  for(let i=comps.length-1;i>0;i--){const j=Math.floor(Math.random()*(i+1));[comps[i],comps[j]]=[comps[j],comps[i]];}
  comps.forEach((comp,i)=>{
    const card=document.createElement('div');card.className='card';card.id='card-'+i;
    card.textContent=comp.name;card.draggable=true;card.dataset.name=comp.name;card.dataset.external=comp.external;
    card.ondragstart=e=>{e.dataTransfer.setData('text/plain',i.toString());card.style.opacity='0.5';};
    card.ondragend=()=>card.style.opacity='1';
    card.ontouchstart=e=>{e.preventDefault();card.style.transform='scale(1.1)';};
    card.ontouchend=e=>{e.preventDefault();card.style.transform='';
      const t=e.changedTouches[0],el=document.elementFromPoint(t.clientX,t.clientY);
      const box=el?el.closest('.sort-box'):null;
      if(box){placeCard(i,box.classList.contains('external')?'external':'internal');}};
    area.appendChild(card);
  });
}
function allowDrop(e){e.preventDefault();e.currentTarget.classList.add('highlight');}
document.querySelectorAll('.sort-box').forEach(b=>{b.ondragleave=()=>b.classList.remove('highlight');});
function dropCard(e,type){e.preventDefault();e.currentTarget.classList.remove('highlight');placeCard(parseInt(e.dataTransfer.getData('text/plain')),type);}
function placeCard(idx,type){
  const comp=LESSON.components[idx];if(!comp)return;
  const card=document.getElementById('card-'+idx);if(card.classList.contains('placed'))return;
  const isCorrect=(type==='external')===comp.external;
  const target=document.getElementById(type==='external'?'extCards':'intCards');
  const placed=document.createElement('div');placed.className='placed-card';placed.textContent=comp.name;
  target.appendChild(placed);card.classList.add('placed');sortCorrect[idx]=isCorrect;
  if(Object.keys(sortCorrect).length===LESSON.components.length){
    const allRight=Object.values(sortCorrect).every(v=>v);
    const fb=document.getElementById('sortFeedback');
    if(allRight){fb.innerHTML='<div class="feedback correct check-anim">Awesome! You sorted all the parts!</div>';
      screenDone[0]=true;document.getElementById('nav1').classList.add('done');celebrate();}
    else{fb.innerHTML='<div class="feedback wrong">Not quite! Try again!</div><button class="reset-btn" onclick="resetSort()">Try Again</button>';}
  }
}
function resetSort(){sortCorrect={};document.getElementById('extCards').innerHTML='';document.getElementById('intCards').innerHTML='';document.getElementById('sortFeedback').innerHTML='';document.querySelectorAll('.card').forEach(c=>c.classList.remove('placed'));}
function initArrows(){
  const model=document.getElementById('arrowModel');let html='';
  LESSON.relationships.forEach((rel,i)=>{
    const wrong=rel.type==='+'?'-':'+';arrowStates[i]=wrong;
    html+='<div class="comp-row"><div class="comp-box">'+rel.from+'</div><div style="text-align:center">'
      +'<button class="arrow-btn '+(wrong==='+'?'positive':'negative')+'" id="arrow-'+i+'" onclick="toggleArrow('+i+')">'
      +(wrong==='+'?'+ &rarr;':'- &rarr;')+'</button><div class="arrow-label">'+(wrong==='+'?'Go together':'Go opposite')
      +'</div></div><div class="comp-box">'+rel.to+'</div></div>';
  });model.innerHTML=html;
}
function toggleArrow(i){
  const next=arrowStates[i]==='+'?'-':'+';arrowStates[i]=next;
  const btn=document.getElementById('arrow-'+i);
  btn.innerHTML=next==='+'?'+ &rarr;':'- &rarr;';
  btn.className='arrow-btn '+(next==='+'?'positive':'negative');
  btn.nextElementSibling.textContent=next==='+'?'Go together':'Go opposite';
  if(LESSON.relationships.every((rel,j)=>arrowStates[j]===rel.type)){
    document.getElementById('arrowFeedback').innerHTML='<div class="feedback correct check-anim">You fixed all the arrows!</div>';
    screenDone[1]=true;document.getElementById('nav2').classList.add('done');celebrate();}
  else document.getElementById('arrowFeedback').innerHTML='';
}
function initExplore(){
  const panel=document.getElementById('explorePanel');let html='';
  const ext=LESSON.components.filter(c=>c.external),int_=LESSON.components.filter(c=>!c.external);
  ext.forEach(comp=>{exploreLevels[comp.name]=50;
    html+='<div class="ctrl-row"><span class="ctrl-label">'+comp.name+'</span>'
      +'<button class="ctrl-btn down" onclick="adj(\''+comp.name+'\',-25)">&#9660;</button>'
      +'<div class="bar-container"><div class="bar-fill ext" id="bar-'+css(comp.name)+'" style="width:50%">50%</div></div>'
      +'<button class="ctrl-btn up" onclick="adj(\''+comp.name+'\',25)">&#9650;</button></div>';});
  int_.forEach(comp=>{exploreLevels[comp.name]=50;
    html+='<div class="ctrl-row"><span class="ctrl-label">'+comp.name+'</span>'
      +'<div class="bar-container"><div class="bar-fill int" id="bar-'+css(comp.name)+'" style="width:50%">50%</div></div></div>';});
  html+='<div class="result-text" id="exploreResult">Press UP or DOWN to explore!</div>';
  html+='<div class="sentence" id="exploreSentence">'+LESSON.sentenceFrame+'</div>';
  panel.innerHTML=html;
}
function css(n){return n.replace(/[^a-zA-Z0-9]/g,'_');}
function adj(name,delta){
  let val=Math.max(0,Math.min(100,exploreLevels[name]+delta));exploreLevels[name]=val;updBar(name,val);
  LESSON.relationships.forEach(rel=>{if(rel.from===name){
    let tv=rel.type==='+'?val:100-val;exploreLevels[rel.to]=tv;updBar(rel.to,tv);}});
  updText();if(!screenDone[2]){screenDone[2]=true;document.getElementById('nav3').classList.add('done');}
}
function updBar(n,v){const b=document.getElementById('bar-'+css(n));if(b){b.style.width=v+'%';b.textContent=v+'%';}}
function updText(){
  const el=document.getElementById('exploreResult'),parts=[];
  LESSON.components.filter(c=>c.external).forEach(comp=>{
    const l=exploreLevels[comp.name],w=l>60?'HIGH':l<40?'LOW':'MEDIUM';
    LESSON.relationships.forEach(rel=>{if(rel.from===comp.name){
      const tl=exploreLevels[rel.to],tw=tl>60?'HIGH':tl<40?'LOW':'MEDIUM';
      parts.push('When '+comp.name+' is '+w+', '+rel.to+' is '+tw+'!');}});});
  el.textContent=parts.join(' ')||'Press buttons to explore!';
}
function celebrate(){
  const div=document.createElement('div');div.className='celebrate';
  const emojis=['&#11088;','&#127775;','&#10024;','&#127881;','&#127882;','&#128079;','&#127942;','&#128171;'];
  for(let i=0;i<24;i++){const s=document.createElement('div');s.className='star';
    s.innerHTML=emojis[Math.floor(Math.random()*emojis.length)];
    s.style.left=Math.random()*100+'%';s.style.animationDelay=Math.random()*1+'s';div.appendChild(s);}
  document.body.appendChild(div);setTimeout(()=>div.remove(),2500);
}
</script>
</body>
</html>'''


def make_game(c, grade_config):
    comps = c.get("components", [])
    rels = c.get("relationships", [])
    lesson_data = {
        "title": c["title"],
        "components": [{"name": n, "desc": d, "external": e} for n, d, e in comps],
        "relationships": [],
        "vocabulary": [{"term": t, "definition": d} for t, d in c.get("vocabulary", [])],
        "sentenceFrame": compose_sentence_frame(c)
    }
    for rel_str, sign, explanation in rels:
        parts = rel_str.split(" to ")
        if len(parts) == 2:
            lesson_data["relationships"].append({
                "from": parts[0].strip(), "to": parts[1].strip(),
                "type": "+" if "+" in sign else "-", "explanation": explanation
            })
    html = GAME_HTML_TEMPLATE.replace("{{LESSON_JSON}}", json.dumps(lesson_data, indent=2))
    return html.replace("{{TITLE}}", c["title"])


def write_game(c, out_dir, grade_config):
    fn = f"{c['id']}-Model-Explorer.html"
    fp = os.path.join(out_dir, fn)
    with open(fp, "w", encoding="utf-8") as f: f.write(make_game(c, grade_config))
    print(f"  [OK] Game: {fn}"); return fp


# ================================================================
#  2. COLORING SHEET (K-2) - LANDSCAPE, FULL-PAGE TPT STYLE
# ================================================================

def make_coloring(c, out_dir, img_dir, grade_label):
    """TPT-style landscape coloring sheet with full-page illustration."""
    lid = c["id"]
    doc = Document()
    make_landscape(doc)

    comps = c.get("components", [])
    rels = c.get("relationships", [])

    # ── Full-page layout table (no visible borders)
    t = doc.add_table(rows=3, cols=1)
    remove_table_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: Header banner
    header_cell = t.rows[0].cells[0]
    cell_shd(header_cell, HEX_LBLUE)
    set_row_height(t.rows[0], 1.1)

    p_title = header_cell.add_paragraph()
    r_t = p_title.add_run("Color My Model!")
    r_t.bold = True; r_t.font.size = DP(32); r_t.font.color.rgb = DMB
    r_t.font.name = "Comic Sans MS"
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = DP(0)

    p_sub = header_cell.add_paragraph()
    r_s = p_sub.add_run(f"{c['title']}  |  {grade_label}")
    r_s.font.size = DP(13); r_s.font.color.rgb = DR(0x44, 0x44, 0x44)
    r_s.font.name = "Comic Sans MS"
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = DP(0)

    name_date_line(header_cell, font_size=12)

    # Row 1: HUGE coloring image (takes most of the page)
    img_cell = t.rows[1].cells[0]
    set_row_height(t.rows[1], 5.5)

    img_path = os.path.join(img_dir, f"{lid}-coloring.png")
    if os.path.exists(img_path):
        p_img = img_cell.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_img = p_img.add_run()
        r_img.add_picture(img_path, width=DI(9.5))
    else:
        p_ph = img_cell.add_paragraph()
        r_ph = p_ph.add_run("[Coloring image will appear here]")
        r_ph.font.size = DP(20); r_ph.font.color.rgb = DR(0xAA, 0xAA, 0xAA)
        p_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 2: Instructions + Color Key (compact footer)
    footer_cell = t.rows[2].cells[0]
    cell_shd(footer_cell, HEX_YELLOW)
    set_row_height(t.rows[2], 1.2)

    # Instructions
    desc = c.get("coloring_description", "")
    if not desc:
        names = [n for n, _, _ in comps]
        desc = f"Color the {', '.join(names[:-1])} and {names[-1]}! Trace the arrows!" if len(names) > 1 else "Color the model!"

    p_inst = footer_cell.add_paragraph()
    r_star = p_inst.add_run("* ")
    r_star.font.size = DP(14); r_star.font.color.rgb = DOG; r_star.bold = True
    r_i = p_inst.add_run(desc)
    r_i.font.size = DP(12); r_i.font.name = "Comic Sans MS"
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.paragraph_format.space_after = DP(2)

    # Color key
    p_key = footer_cell.add_paragraph()
    r_k = p_key.add_run("Arrow Color Key:  ")
    r_k.font.size = DP(11); r_k.bold = True; r_k.font.name = "Comic Sans MS"
    r_g = p_key.add_run("GREEN (+) = Go together   ")
    r_g.font.size = DP(11); r_g.font.color.rgb = DGN; r_g.bold = True; r_g.font.name = "Comic Sans MS"
    r_r = p_key.add_run("RED (-) = Go opposite")
    r_r.font.size = DP(11); r_r.font.color.rgb = DRD; r_r.bold = True; r_r.font.name = "Comic Sans MS"
    p_key.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_brand = footer_cell.add_paragraph()
    r_b = p_brand.add_run(f"ModelIt! {lid}")
    r_b.font.size = DP(8); r_b.font.color.rgb = DR(0x99, 0x99, 0x99)
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fn = f"{lid}-Coloring-Sheet.docx"
    out = os.path.join(out_dir, fn)
    doc.save(out)
    print(f"  [OK] Coloring Sheet: {fn}")
    return out


# ================================================================
#  3. FIX THE SYSTEM (G2-3) - LANDSCAPE, VISUAL
# ================================================================

def make_fix_activity(c, out_dir, img_dir, grade_label):
    """Landscape fix-the-system worksheet with visual diagrams."""
    lid = c["id"]
    doc = Document()
    make_landscape(doc)

    comps = c.get("components", [])
    rels = c.get("relationships", [])
    vocab = c.get("vocabulary", [])

    # ── PAGE 1: Can You Fix It? ──

    # Header banner
    t_hdr = doc.add_table(rows=1, cols=1)
    remove_table_borders(t_hdr)
    cell_h = t_hdr.rows[0].cells[0]
    cell_shd(cell_h, HEX_LORANGE)
    set_row_height(t_hdr.rows[0], 0.9)

    fun_header(cell_h, "Can You Fix It?", f"{c['title']}  |  {grade_label}", font_size=26, color=DOG)
    name_date_line(cell_h, font_size=11)

    # Optional scene image
    img_path = os.path.join(img_dir, f"{lid}-fix-diagram.png")
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_img = p_img.add_run()
        r_img.add_picture(img_path, width=DI(4.5))
        p_img.paragraph_format.space_after = DP(4)

    # 3-column layout for the 3 errors
    t_fix = doc.add_table(rows=1, cols=3)
    t_fix.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ── Error 1: Wrong Direction
    c1 = t_fix.rows[0].cells[0]
    cell_shd(c1, HEX_LPINK)
    set_cell_border(c1, top=("sz", "12", "color", HEX_RED, "val", "single"),
                    bottom=("sz", "12", "color", HEX_RED, "val", "single"),
                    left=("sz", "12", "color", HEX_RED, "val", "single"),
                    right=("sz", "12", "color", HEX_RED, "val", "single"))

    p1h = c1.add_paragraph()
    r1h = p1h.add_run("1. Wrong Direction!")
    r1h.bold = True; r1h.font.size = DP(13); r1h.font.color.rgb = DRD
    r1h.font.name = "Comic Sans MS"
    p1h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if rels:
        rel_str, sign, _ = rels[0]
        parts = rel_str.split(" to ")
        fr = parts[0].strip() if len(parts) >= 2 else "Part A"
        to = parts[1].strip() if len(parts) >= 2 else "Part B"
        wrong_sign = "-" if "+" in sign else "+"

        p1d = c1.add_paragraph()
        p1d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_fr = p1d.add_run(f"{fr}")
        r_fr.bold = True; r_fr.font.size = DP(12); r_fr.font.name = "Comic Sans MS"
        r_arr = p1d.add_run(f"  {wrong_sign} -->  ")
        r_arr.font.size = DP(14); r_arr.bold = True
        r_arr.font.color.rgb = DRD if wrong_sign == "-" else DGN
        r_to = p1d.add_run(f"{to}")
        r_to.bold = True; r_to.font.size = DP(12); r_to.font.name = "Comic Sans MS"

    p1q = c1.add_paragraph()
    r1q = p1q.add_run("Circle the wrong sign! Write the correct one:")
    r1q.font.size = DP(10); r1q.font.name = "Comic Sans MS"
    p1q.paragraph_format.space_before = DP(6)
    writing_lines(c1, count=2, width=25)

    # ── Error 2: Wrong Category
    c2 = t_fix.rows[0].cells[1]
    cell_shd(c2, HEX_LPURP)
    set_cell_border(c2, top=("sz", "12", "color", "9B59B6", "val", "single"),
                    bottom=("sz", "12", "color", "9B59B6", "val", "single"),
                    left=("sz", "12", "color", "9B59B6", "val", "single"),
                    right=("sz", "12", "color", "9B59B6", "val", "single"))

    p2h = c2.add_paragraph()
    r2h = p2h.add_run("2. Wrong Category!")
    r2h.bold = True; r2h.font.size = DP(13); r2h.font.color.rgb = DPURP
    r2h.font.name = "Comic Sans MS"
    p2h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(comps) >= 2:
        swap_name, swap_desc, swap_ext = comps[0]
        wrong_cat = "Inside" if swap_ext else "Outside"

        for idx, (name, desc, is_ext) in enumerate(comps):
            p_c = c2.add_paragraph()
            cat = wrong_cat if idx == 0 else ("Outside" if is_ext else "Inside")
            r_n = p_c.add_run(f"{name} = {cat}")
            r_n.font.size = DP(11); r_n.font.name = "Comic Sans MS"
            r_n.bold = True if idx == 0 else False
            if idx == 0:
                r_n.font.color.rgb = DPURP
            p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2q = c2.add_paragraph()
    r2q = p2q.add_run("Which part is in the wrong place? Fix it!")
    r2q.font.size = DP(10); r2q.font.name = "Comic Sans MS"
    p2q.paragraph_format.space_before = DP(6)
    writing_lines(c2, count=2, width=25)

    # ── Error 3: Missing Part
    c3 = t_fix.rows[0].cells[2]
    cell_shd(c3, HEX_LGREEN)
    set_cell_border(c3, top=("sz", "12", "color", HEX_GREEN, "val", "single"),
                    bottom=("sz", "12", "color", HEX_GREEN, "val", "single"),
                    left=("sz", "12", "color", HEX_GREEN, "val", "single"),
                    right=("sz", "12", "color", HEX_GREEN, "val", "single"))

    p3h = c3.add_paragraph()
    r3h = p3h.add_run("3. Missing Part!")
    r3h.bold = True; r3h.font.size = DP(13); r3h.font.color.rgb = DGN
    r3h.font.name = "Comic Sans MS"
    p3h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(comps) >= 2:
        missing_idx = len(comps) - 1
        missing_name = comps[missing_idx][0]

        for idx, (name, desc, is_ext) in enumerate(comps):
            p_m = c3.add_paragraph()
            display = "???" if idx == missing_idx else name
            r_m = p_m.add_run(display)
            r_m.font.size = DP(12); r_m.font.name = "Comic Sans MS"; r_m.bold = True
            r_m.font.color.rgb = DRD if idx == missing_idx else DN
            p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Word bank
    p_wb_h = c3.add_paragraph()
    r_wbh = p_wb_h.add_run("Word Bank:")
    r_wbh.font.size = DP(9); r_wbh.bold = True; r_wbh.font.name = "Comic Sans MS"
    p_wb_h.paragraph_format.space_before = DP(4)

    comp_names = [n for n, _, _ in comps]
    distractors = [t for t, d in vocab if t not in comp_names][:1]
    bank = comp_names + distractors
    random.shuffle(bank)

    p_wb = c3.add_paragraph()
    r_wb = p_wb.add_run("  |  ".join(bank))
    r_wb.font.size = DP(9); r_wb.font.name = "Comic Sans MS"; r_wb.bold = True
    p_wb.alignment = WD_ALIGN_PARAGRAPH.CENTER

    writing_lines(c3, count=2, width=25)

    # ── PAGE 2: Build Your Own ──
    doc.add_page_break()

    t_hdr2 = doc.add_table(rows=1, cols=1)
    remove_table_borders(t_hdr2)
    cell_h2 = t_hdr2.rows[0].cells[0]
    cell_shd(cell_h2, HEX_LBLUE)
    set_row_height(t_hdr2.rows[0], 0.7)
    fun_header(cell_h2, "Build Your Own Model!", c['title'], font_size=24, color=DMB)

    # Big drawing area
    t_draw = doc.add_table(rows=1, cols=1)
    t_draw.style = "Table Grid"
    t_draw.alignment = WD_TABLE_ALIGNMENT.CENTER
    draw_cell = t_draw.rows[0].cells[0]
    set_row_height(t_draw.rows[0], 4.5)

    p_inst = draw_cell.add_paragraph()
    r_inst = p_inst.add_run("Draw the parts in boxes. Draw arrows to connect them. Write + or - on each arrow!")
    r_inst.font.size = DP(13); r_inst.font.name = "Comic Sans MS"; r_inst.font.color.rgb = DR(0x66, 0x66, 0x66)
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sentence frame
    doc.add_paragraph()
    frame = compose_sentence_frame(c)
    p_frame = doc.add_paragraph()
    r_f = p_frame.add_run(f"Complete: {frame}")
    r_f.font.size = DP(16); r_f.bold = True; r_f.font.color.rgb = DBB
    r_f.font.name = "Comic Sans MS"
    p_frame.alignment = WD_ALIGN_PARAGRAPH.CENTER

    writing_lines(doc, count=2, width=55)

    # Footer
    p_ft = doc.add_paragraph()
    r_ft = p_ft.add_run(f"ModelIt! {lid}")
    r_ft.font.size = DP(8); r_ft.font.color.rgb = DR(0x99, 0x99, 0x99)
    p_ft.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fn = f"{lid}-Fix-the-System.docx"
    out = os.path.join(out_dir, fn)
    doc.save(out)
    print(f"  [OK] Fix the System: {fn}")
    return out


# ================================================================
#  4. INTERACTIVE NOTEBOOK INSERT (G2-3) - LANDSCAPE, FOLDABLE
# ================================================================

def make_notebook(c, out_dir, grade_label):
    """Landscape foldable interactive notebook insert - visual and colorful."""
    lid = c["id"]
    doc = Document()
    make_landscape(doc)

    comps = c.get("components", [])
    rels = c.get("relationships", [])
    vocab = c.get("vocabulary", [])
    frame = compose_sentence_frame(c)

    # Cut line top
    p_cut = doc.add_paragraph()
    r_cut = p_cut.add_run("- - - - - - - - - - - - - - - - - - CUT HERE - - - - - - - - - - - - - - - - - -")
    r_cut.font.size = DP(9); r_cut.font.color.rgb = DR(0xAA, 0xAA, 0xAA)
    r_cut.font.name = "Comic Sans MS"
    p_cut.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Main 2-column table
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ── LEFT: Front (visible when glued in) ──
    front = t.rows[0].cells[0]
    cell_shd(front, HEX_LBLUE)
    set_cell_width(front, 4.9)
    set_cell_border(front, top=("sz", "18", "color", HEX_MB, "val", "single"),
                    bottom=("sz", "18", "color", HEX_MB, "val", "single"),
                    left=("sz", "18", "color", HEX_MB, "val", "single"),
                    right=("sz", "6", "color", HEX_MB, "val", "dashed"))

    fun_header(front, "My System Model", c["title"], font_size=18, color=DMB)

    p_grade = front.add_paragraph()
    r_gr = p_grade.add_run(f"{grade_label}  |  {c.get('ngss', '')}")
    r_gr.font.size = DP(8); r_gr.font.color.rgb = DR(0x55, 0x55, 0x55)
    r_gr.font.name = "Comic Sans MS"
    p_grade.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Component diagram
    p_mh = front.add_paragraph()
    r_mh = p_mh.add_run("Color the model!")
    r_mh.bold = True; r_mh.font.size = DP(11); r_mh.font.color.rgb = DOG
    r_mh.font.name = "Comic Sans MS"
    p_mh.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if comps:
        inner = front.add_table(rows=2, cols=len(comps))
        remove_table_borders(inner)
        for idx, (name, desc, is_ext) in enumerate(comps):
            # Label row
            lbl_cell = inner.rows[0].cells[idx]
            p_lbl = lbl_cell.add_paragraph()
            tag = "OUTSIDE" if is_ext else "INSIDE"
            r_lbl = p_lbl.add_run(tag)
            r_lbl.font.size = DP(7); r_lbl.bold = True
            r_lbl.font.color.rgb = DOG if is_ext else DMB
            r_lbl.font.name = "Comic Sans MS"
            p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Name box row
            name_cell = inner.rows[1].cells[idx]
            hex_bg = HEX_LORANGE if is_ext else HEX_LGREEN
            cell_shd(name_cell, hex_bg)
            set_cell_border(name_cell, top=("sz", "8", "color", HEX_BB, "val", "single"),
                            bottom=("sz", "8", "color", HEX_BB, "val", "single"),
                            left=("sz", "8", "color", HEX_BB, "val", "single"),
                            right=("sz", "8", "color", HEX_BB, "val", "single"))
            p_nm = name_cell.add_paragraph()
            r_nm = p_nm.add_run(name)
            r_nm.font.size = DP(11); r_nm.bold = True; r_nm.font.color.rgb = DN
            r_nm.font.name = "Comic Sans MS"
            p_nm.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Arrow legend
    p_key = front.add_paragraph()
    r_k1 = p_key.add_run("GREEN + = go together  ")
    r_k1.font.size = DP(9); r_k1.font.color.rgb = DGN; r_k1.bold = True
    r_k1.font.name = "Comic Sans MS"
    r_k2 = p_key.add_run("RED - = go opposite")
    r_k2.font.size = DP(9); r_k2.font.color.rgb = DRD; r_k2.bold = True
    r_k2.font.name = "Comic Sans MS"
    p_key.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fold indicator
    p_fold = front.add_paragraph()
    r_fold = p_fold.add_run("--- FOLD HERE ---")
    r_fold.font.size = DP(8); r_fold.font.color.rgb = DR(0x99, 0x99, 0x99)
    r_fold.font.name = "Comic Sans MS"
    p_fold.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── RIGHT: Inside (hidden when folded) ──
    inside = t.rows[0].cells[1]
    cell_shd(inside, HEX_YELLOW)
    set_cell_width(inside, 4.9)
    set_cell_border(inside, top=("sz", "18", "color", HEX_OG, "val", "single"),
                    bottom=("sz", "18", "color", HEX_OG, "val", "single"),
                    left=("sz", "6", "color", HEX_OG, "val", "dashed"),
                    right=("sz", "18", "color", HEX_OG, "val", "single"))

    fun_header(inside, "What I Learned", "", font_size=16, color=DOG)

    # Vocabulary
    p_vh = inside.add_paragraph()
    r_vh = p_vh.add_run("Key Words:")
    r_vh.bold = True; r_vh.font.size = DP(11); r_vh.font.color.rgb = DN
    r_vh.font.name = "Comic Sans MS"

    for term, defn in vocab[:4]:
        p_v = inside.add_paragraph()
        r_t = p_v.add_run(f"{term}: ")
        r_t.bold = True; r_t.font.size = DP(9); r_t.font.color.rgb = DBB
        r_t.font.name = "Comic Sans MS"
        r_d = p_v.add_run(defn)
        r_d.font.size = DP(8); r_d.font.name = "Comic Sans MS"

    # Sentence frame
    p_sfh = inside.add_paragraph()
    r_sfh = p_sfh.add_run("Complete:")
    r_sfh.bold = True; r_sfh.font.size = DP(10); r_sfh.font.color.rgb = DN
    r_sfh.font.name = "Comic Sans MS"
    p_sfh.paragraph_format.space_before = DP(6)

    p_sf = inside.add_paragraph()
    r_sf = p_sf.add_run(frame)
    r_sf.font.size = DP(10); r_sf.bold = True; r_sf.font.color.rgb = DBB
    r_sf.font.name = "Comic Sans MS"

    # Writing
    p_wh = inside.add_paragraph()
    r_wh = p_wh.add_run("I learned that...")
    r_wh.bold = True; r_wh.font.size = DP(10); r_wh.font.color.rgb = DN
    r_wh.font.name = "Comic Sans MS"
    p_wh.paragraph_format.space_before = DP(6)
    writing_lines(inside, count=3, width=40)

    # Glue tab
    p_glue = inside.add_paragraph()
    r_glue = p_glue.add_run("-- GLUE HERE --")
    r_glue.font.size = DP(7); r_glue.font.color.rgb = DR(0xAA, 0xAA, 0xAA)
    r_glue.font.name = "Comic Sans MS"
    p_glue.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Cut line bottom
    p_cut2 = doc.add_paragraph()
    r_cut2 = p_cut2.add_run("- - - - - - - - - - - - - - - - - - CUT HERE - - - - - - - - - - - - - - - - - -")
    r_cut2.font.size = DP(9); r_cut2.font.color.rgb = DR(0xAA, 0xAA, 0xAA)
    r_cut2.font.name = "Comic Sans MS"
    p_cut2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fn = f"{lid}-Interactive-Notebook.docx"
    out = os.path.join(out_dir, fn)
    doc.save(out)
    print(f"  [OK] Interactive Notebook: {fn}")
    return out


# ================================================================
#  5. READING PASSAGE (G2-3) - PICTURE-DRIVEN
# ================================================================

def compose_passage(c):
    """Build 150-200 word passage from lesson data."""
    big_q = c.get("big_question", c.get("title", ""))
    bg_sections = c.get("background_sections", [])
    bg_intro = c.get("background_intro", "")
    answer = c.get("answer", "")

    hook = f"Have you ever wondered: {big_q} Let's find out!"
    body1 = ""
    if bg_sections:
        sentences = bg_sections[0][1].split(".")
        body1 = ". ".join(s.strip() for s in sentences[:2] if s.strip()) + "."
    elif bg_intro:
        sentences = bg_intro.split(".")
        body1 = ". ".join(s.strip() for s in sentences[:2] if s.strip()) + "."

    body2 = ""
    if len(bg_sections) >= 2:
        sentences = bg_sections[1][1].split(".")
        body2 = ". ".join(s.strip() for s in sentences[:2] if s.strip()) + "."

    closing = ""
    if answer:
        sentences = answer.split(".")
        closing = "Now you know! " + ". ".join(s.strip() for s in sentences[:2] if s.strip()) + "."

    passage = f"{hook}\n\n{body1}"
    if body2: passage += f"\n\n{body2}"
    passage += f"\n\n{closing}"

    words = passage.split()
    if len(words) > 220:
        passage = " ".join(words[:200]) + "..."
    return passage


def make_reading(c, out_dir, img_dir, grade_label):
    """Picture-driven reading passage with scene illustration."""
    lid = c["id"]
    doc = Document()
    make_portrait_tight(doc)

    comps = c.get("components", [])
    vocab = c.get("vocabulary", [])
    misconceptions = c.get("misconceptions", [])
    big_q = c.get("big_question", c["title"])

    # ── Header banner
    t_hdr = doc.add_table(rows=1, cols=1)
    remove_table_borders(t_hdr)
    cell_h = t_hdr.rows[0].cells[0]
    cell_shd(cell_h, HEX_LBLUE)
    fun_header(cell_h, big_q, "", font_size=22, color=DMB)

    p_sub = cell_h.add_paragraph()
    r_sub = p_sub.add_run(f"ModelIt! Reading  |  {grade_label}")
    r_sub.font.size = DP(10); r_sub.font.color.rgb = DR(0x55, 0x55, 0x55)
    r_sub.font.name = "Comic Sans MS"
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    name_date_line(cell_h, font_size=11)

    # ── Scene image (if generated)
    img_path = os.path.join(img_dir, f"{lid}-scene.png")
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_img = p_img.add_run()
        r_img.add_picture(img_path, width=DI(5.5))
        p_img.paragraph_format.space_after = DP(6)

    # ── Reading passage
    passage = compose_passage(c)
    for para_text in passage.split("\n\n"):
        if para_text.strip():
            p = doc.add_paragraph()
            r = p.add_run(para_text.strip())
            r.font.size = DP(13); r.font.name = "Comic Sans MS"
            p.paragraph_format.space_after = DP(6)
            p.paragraph_format.line_spacing = DP(22)

    # ── Fun Fact box
    if misconceptions:
        misc_text, reality, _ = misconceptions[0]
        t_fun = doc.add_table(rows=1, cols=1)
        t_fun.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_fun = t_fun.rows[0].cells[0]
        cell_shd(cell_fun, HEX_YELLOW)
        set_cell_border(cell_fun, top=("sz", "12", "color", HEX_OG, "val", "single"),
                        bottom=("sz", "12", "color", HEX_OG, "val", "single"),
                        left=("sz", "12", "color", HEX_OG, "val", "single"),
                        right=("sz", "12", "color", HEX_OG, "val", "single"))

        p_fh = cell_fun.add_paragraph()
        r_fh = p_fh.add_run("Fun Fact!")
        r_fh.bold = True; r_fh.font.size = DP(16); r_fh.font.color.rgb = DOG
        r_fh.font.name = "Comic Sans MS"
        p_fh.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p_fb = cell_fun.add_paragraph()
        r_fb = p_fb.add_run(f"Some people think: \"{misc_text}\" -- but that's not true! {reality}")
        r_fb.font.size = DP(11); r_fb.font.name = "Comic Sans MS"
        p_fb.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Think About It
    t_think = doc.add_table(rows=1, cols=1)
    t_think.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_think = t_think.rows[0].cells[0]
    cell_shd(cell_think, HEX_LGREEN)
    set_cell_border(cell_think, top=("sz", "12", "color", HEX_GREEN, "val", "single"),
                    bottom=("sz", "12", "color", HEX_GREEN, "val", "single"),
                    left=("sz", "12", "color", HEX_GREEN, "val", "single"),
                    right=("sz", "12", "color", HEX_GREEN, "val", "single"))

    p_th = cell_think.add_paragraph()
    r_th = p_th.add_run("Think About It")
    r_th.bold = True; r_th.font.size = DP(14); r_th.font.color.rgb = DGN
    r_th.font.name = "Comic Sans MS"
    p_th.alignment = WD_ALIGN_PARAGRAPH.CENTER

    questions = []
    if big_q:
        questions.append(f"What surprised you most about {big_q.lower().rstrip('?')}?")
    questions.append("How could you explain what you learned to a friend or family member?")

    for i, q in enumerate(questions[:2], 1):
        p_q = cell_think.add_paragraph()
        r_q = p_q.add_run(f"{i}. {q}")
        r_q.font.size = DP(12); r_q.bold = True; r_q.font.name = "Comic Sans MS"
        writing_lines(cell_think, count=3, width=55)

    # Footer
    p_ft = doc.add_paragraph()
    r_ft = p_ft.add_run(f"ModelIt! {lid}")
    r_ft.font.size = DP(8); r_ft.font.color.rgb = DR(0x99, 0x99, 0x99)
    p_ft.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fn = f"{lid}-Reading-Passage.docx"
    out = os.path.join(out_dir, fn)
    doc.save(out)
    print(f"  [OK] Reading Passage: {fn}")
    return out


# ================================================================
#  6. PARENT-HOME CONNECTION (G2-3) - COLORFUL SINGLE PAGE
# ================================================================

def make_home_connection(c, out_dir, grade_label):
    """Colorful, visual parent-home connection sheet."""
    lid = c["id"]
    doc = Document()
    make_portrait_tight(doc)

    vocab = c.get("vocabulary", [])
    scenarios = c.get("scenarios", [])
    cross_curr = c.get("cross_curr", [])
    big_q = c.get("big_question", c["title"])
    frame = compose_sentence_frame(c)

    # ── Header
    t_hdr = doc.add_table(rows=1, cols=1)
    remove_table_borders(t_hdr)
    cell_h = t_hdr.rows[0].cells[0]
    cell_shd(cell_h, HEX_LBLUE)
    set_row_height(t_hdr.rows[0], 0.9)
    fun_header(cell_h, "Family Science Connection", f"{grade_label}  |  {lid}", font_size=22, color=DMB)

    # Dear Family
    p_dear = doc.add_paragraph()
    r_d = p_dear.add_run(f"Dear Family, today your child learned about: {c['title']}!")
    r_d.font.size = DP(13); r_d.italic = True; r_d.font.name = "Comic Sans MS"
    p_dear.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dear.paragraph_format.space_after = DP(6)

    # ── 4 sections in a 2x2 grid
    t_grid = doc.add_table(rows=2, cols=2)
    t_grid.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Section 1: Talk About It (top-left)
    c1 = t_grid.rows[0].cells[0]
    cell_shd(c1, HEX_LORANGE)
    set_cell_border(c1, top=("sz", "12", "color", HEX_OG, "val", "single"),
                    bottom=("sz", "12", "color", HEX_OG, "val", "single"),
                    left=("sz", "12", "color", HEX_OG, "val", "single"),
                    right=("sz", "12", "color", HEX_OG, "val", "single"))

    p1h = c1.add_paragraph()
    r1h = p1h.add_run("Talk About It")
    r1h.bold = True; r1h.font.size = DP(14); r1h.font.color.rgb = DOG
    r1h.font.name = "Comic Sans MS"
    p1h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1b = c1.add_paragraph()
    r1b = p1b.add_run(f"Ask your child: \"{big_q}\" and listen to what they say!")
    r1b.font.size = DP(10); r1b.font.name = "Comic Sans MS"

    # Section 2: Try It at Home (top-right)
    c2 = t_grid.rows[0].cells[1]
    cell_shd(c2, HEX_LGREEN)
    set_cell_border(c2, top=("sz", "12", "color", HEX_GREEN, "val", "single"),
                    bottom=("sz", "12", "color", HEX_GREEN, "val", "single"),
                    left=("sz", "12", "color", HEX_GREEN, "val", "single"),
                    right=("sz", "12", "color", HEX_GREEN, "val", "single"))

    p2h = c2.add_paragraph()
    r2h = p2h.add_run("Try It at Home")
    r2h.bold = True; r2h.font.size = DP(14); r2h.font.color.rgb = DGN
    r2h.font.name = "Comic Sans MS"
    p2h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if scenarios:
        p2b = c2.add_paragraph()
        r2b = p2b.add_run(f"No special materials needed!\n{scenarios[0][1] if len(scenarios[0]) > 1 else scenarios[0][0]}")
        r2b.font.size = DP(10); r2b.font.name = "Comic Sans MS"
    else:
        p2b = c2.add_paragraph()
        r2b = p2b.add_run("Observe this topic together at home or on a walk!")
        r2b.font.size = DP(10); r2b.font.name = "Comic Sans MS"

    # Section 3: Words to Know (bottom-left)
    c3 = t_grid.rows[1].cells[0]
    cell_shd(c3, HEX_YELLOW)
    set_cell_border(c3, top=("sz", "12", "color", "F39C12", "val", "single"),
                    bottom=("sz", "12", "color", "F39C12", "val", "single"),
                    left=("sz", "12", "color", "F39C12", "val", "single"),
                    right=("sz", "12", "color", "F39C12", "val", "single"))

    p3h = c3.add_paragraph()
    r3h = p3h.add_run("Words to Know")
    r3h.bold = True; r3h.font.size = DP(14); r3h.font.color.rgb = DYEL
    r3h.font.name = "Comic Sans MS"
    p3h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for term, defn in vocab[:4]:
        p_v = c3.add_paragraph()
        r_t = p_v.add_run(f"{term}: ")
        r_t.bold = True; r_t.font.size = DP(10); r_t.font.color.rgb = DBB
        r_t.font.name = "Comic Sans MS"
        r_def = p_v.add_run(defn)
        r_def.font.size = DP(10); r_def.font.name = "Comic Sans MS"

    # Section 4: Keep Exploring (bottom-right)
    c4 = t_grid.rows[1].cells[1]
    cell_shd(c4, HEX_LPURP)
    set_cell_border(c4, top=("sz", "12", "color", "9B59B6", "val", "single"),
                    bottom=("sz", "12", "color", "9B59B6", "val", "single"),
                    left=("sz", "12", "color", "9B59B6", "val", "single"),
                    right=("sz", "12", "color", "9B59B6", "val", "single"))

    p4h = c4.add_paragraph()
    r4h = p4h.add_run("Keep Exploring")
    r4h.bold = True; r4h.font.size = DP(14); r4h.font.color.rgb = DPURP
    r4h.font.name = "Comic Sans MS"
    p4h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if cross_curr:
        for subj, conn in cross_curr[:2]:
            p_cc = c4.add_paragraph()
            r_s = p_cc.add_run(f"{subj}: ")
            r_s.bold = True; r_s.font.size = DP(10); r_s.font.name = "Comic Sans MS"
            r_c = p_cc.add_run(conn)
            r_c.font.size = DP(10); r_c.font.name = "Comic Sans MS"
    else:
        p_cc = c4.add_paragraph()
        r_cc = p_cc.add_run("Go on a nature walk and look for examples of what we learned!")
        r_cc.font.size = DP(10); r_cc.font.name = "Comic Sans MS"

    doc.add_paragraph()

    # Footer with sentence frame
    t_ft = doc.add_table(rows=1, cols=1)
    remove_table_borders(t_ft)
    cell_ft = t_ft.rows[0].cells[0]
    cell_shd(cell_ft, HEX_LBLUE)

    p_ask = cell_ft.add_paragraph()
    r_ask = p_ask.add_run(f"Ask your child: \"{frame}\"")
    r_ask.font.size = DP(12); r_ask.bold = True; r_ask.font.color.rgb = DBB
    r_ask.font.name = "Comic Sans MS"
    p_ask.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_brand = cell_ft.add_paragraph()
    r_brand = p_brand.add_run(f"ModelIt! {lid}")
    r_brand.font.size = DP(8); r_brand.font.color.rgb = DR(0x99, 0x99, 0x99)
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fn = f"{lid}-Parent-Home-Connection.docx"
    out = os.path.join(out_dir, fn)
    doc.save(out)
    print(f"  [OK] Parent-Home Connection: {fn}")
    return out


# ================================================================
#  MAIN ORCHESTRATOR
# ================================================================

def load_lessons(grade_key):
    cfg = GRADE_CONFIG[grade_key]
    mod = importlib.import_module(cfg["data"])
    return mod.ALL_LESSONS


def generate_supplementary(c, grade_key):
    cfg = GRADE_CONFIG[grade_key]
    lid = c["id"]
    base_dir = os.path.join(SCRIPT_DIR, '..', 'materials', cfg["dir"])
    out_dir = os.path.join(base_dir, lid)
    img_dir = os.path.join(out_dir, "images")
    os.makedirs(out_dir, exist_ok=True)
    os.makedirs(img_dir, exist_ok=True)
    grade_label = cfg["label"]
    count = 0

    print(f"\n{'-'*50}")
    print(f"  {lid}: {c['title']} ({grade_label})")
    print(f"{'-'*50}")

    # K-2: Game
    if cfg["game"]:
        write_game(c, out_dir, cfg)
        count += 1

    # K-2: Coloring (landscape, full-page TPT style, AI image)
    if cfg["coloring"]:
        gen_coloring_img(c, img_dir)
        make_coloring(c, out_dir, img_dir, grade_label)
        count += 1
        time.sleep(3)

    # G2-3: Fix the System (landscape, visual, AI diagram)
    if cfg["fix"]:
        gen_fix_img(c, img_dir)
        make_fix_activity(c, out_dir, img_dir, grade_label)
        count += 1
        time.sleep(3)

    # G2-3: Interactive Notebook (landscape, colorful)
    if cfg["notebook"]:
        make_notebook(c, out_dir, grade_label)
        count += 1

    # G2-3: Reading Passage (picture-driven, scene image)
    if cfg["reading"]:
        gen_scene_img(c, img_dir, "scene")
        make_reading(c, out_dir, img_dir, grade_label)
        count += 1
        time.sleep(3)

    # G2-3: Parent-Home Connection (colorful grid)
    if cfg["home"]:
        make_home_connection(c, out_dir, grade_label)
        count += 1

    print(f"  [{count} files for {lid}]")
    return count


def run_grade(grade_key):
    cfg = GRADE_CONFIG[grade_key]
    lessons = load_lessons(grade_key)

    types = []
    if cfg["game"]: types.append("Game")
    if cfg["coloring"]: types.append("Coloring")
    if cfg["fix"]: types.append("Fix Activity")
    if cfg["notebook"]: types.append("Notebook")
    if cfg["reading"]: types.append("Reading")
    if cfg["home"]: types.append("Home Connection")

    print("\n" + "=" * 60)
    print(f"  ModelIt! Supplementary Materials (TPT Edition) - {cfg['label']}")
    print(f"  {len(lessons)} lessons x {len(types)} types = {len(lessons) * len(types)} files")
    print(f"  Types: {', '.join(types)}")
    print("=" * 60)

    for lesson in lessons:
        print(f"  - {lesson['id']}: {lesson['title']}")
    print()

    start = time.time()
    total = 0
    for i, lesson in enumerate(lessons, 1):
        print(f"\n[{i}/{len(lessons)}] Processing {lesson['id']}...")
        total += generate_supplementary(lesson, grade_key)

    elapsed = time.time() - start
    print("\n" + "=" * 60)
    print(f"  BATCH COMPLETE: {cfg['label']}")
    print(f"  Files: {total}  |  Time: {elapsed/60:.1f}m  |  Image cost: ${total_cost:.2f}")
    print("=" * 60)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python create_supplementary_K3.py [K|1|2|3|all]")
        print("  K/1: game + coloring (landscape, TPT-style)")
        print("  2: all 6 types (game+coloring+fix+notebook+reading+home)")
        print("  3: fix+notebook+reading+home")
        sys.exit(1)

    grade = sys.argv[1].upper()
    if grade == "ALL":
        for g in ["K", "1", "2", "3"]: run_grade(g)
    elif grade in GRADE_CONFIG:
        run_grade(grade)
    else:
        print(f"Unknown grade: {grade}. Use K, 1, 2, 3, or all.")
        sys.exit(1)
